"""
Celery tasks for asynchronous survey generation.

Key changes from original:
  1. Removed autoretry_for / retry_backoff decorator args — manual self.retry()
     is the sole retry mechanism to prevent double-firing.
  2. R2 upload failure now falls back to local disk instead of silently
     marking the survey COMPLETED with a broken download link.
  3. Survey is only marked COMPLETED after a usable doc_link is confirmed.
"""
import json
import time
from io import BytesIO
from pathlib import Path
from typing import Dict, Any

import asyncio
from docx import Document

from app.core.celery import celery_app
from app.core.logging import get_logger
from app.models.database import SessionLocal
from app.models.survey import SurveyRequestRecord
from app.services.ai_service import AIService
import redis.asyncio as aioredis
from app.core.config import settings

logger = get_logger(__name__)


def _parse_questions_payload(raw: str | dict) -> dict[str, Any]:
    if isinstance(raw, dict):
        return raw

    if not isinstance(raw, str):
        raise TypeError(f"Unsupported questions payload type: {type(raw)!r}")

    candidates = [raw.strip()]
    trimmed = raw.strip()

    if trimmed.startswith("```"):
        fenced = trimmed.split("\n")
        if fenced and fenced[0].startswith("```"):
            fenced = fenced[1:]
        if fenced and fenced[-1].strip().startswith("```"):
            fenced = fenced[:-1]
        candidates.append("\n".join(fenced).strip())

    start_idx = trimmed.find("{")
    end_idx = trimmed.rfind("}")
    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        candidates.append(trimmed[start_idx:end_idx + 1])

    for candidate in candidates:
        if not candidate:
            continue
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            continue

    logger.warning("questions_json_parse_failed", preview=trimmed[:240])
    return {
        "questions": [
            {
                "text": "What best describes your current relationship with this company?",
                "type": "radiogroup",
                "choices": ["Existing customer", "Prospect", "Former customer", "Not sure"],
            },
            {
                "text": "How familiar are you with the company's offerings?",
                "type": "radiogroup",
                "choices": ["Very familiar", "Somewhat familiar", "Not very familiar", "Not at all familiar"],
            },
            {
                "text": "What is the main reason you would engage with this company?",
                "type": "comment",
                "choices": [],
            },
            {
                "text": "How likely are you to recommend this company to others?",
                "type": "radiogroup",
                "choices": ["Very likely", "Somewhat likely", "Neutral", "Somewhat unlikely", "Very unlikely"],
            },
            {
                "text": "What could the company improve to better meet your needs?",
                "type": "comment",
                "choices": [],
            },
        ]
    }


# ── Redis progress helper ─────────────────────────────────────────────────────

async def publish_progress(request_id: str, message: str) -> None:
    try:
        r = await asyncio.wait_for(
            aioredis.from_url(settings.REDIS_URL, decode_responses=True),
            timeout=2.0,
        )
        await r.publish(f"survey_progress_{request_id}", message)
        await r.close()
    except (asyncio.TimeoutError, ConnectionError) as exc:
        logger.debug("redis_publish_skipped", request_id=request_id, error=str(exc))


# ── DB status helper ──────────────────────────────────────────────────────────

def update_survey_status(
    request_id: str,
    status: str,
    pages=None,
    questionnaire_data=None,
    doc_link: str | None = None,
    business_overview: str | None = None,
    use_case: str | None = None,
    research_objectives: str | None = None,
) -> None:
    db = SessionLocal()
    try:
        record = (
            db.query(SurveyRequestRecord)
            .filter(SurveyRequestRecord.request_id == request_id)
            .first()
        )
        if record:
            record.status = status
            if pages is not None:
                record.pages = pages
            if questionnaire_data is not None:
                record.questionnaire_data = questionnaire_data
            if doc_link is not None:
                record.doc_link = doc_link
            if business_overview is not None:
                record.business_overview = business_overview
            if use_case is not None:
                record.use_case = use_case
            if research_objectives is not None:
                record.research_objectives = research_objectives
            db.commit()
            logger.info("survey_status_updated", request_id=request_id, status=status)
    except Exception as exc:
        logger.error("survey_status_update_failed", request_id=request_id, error=str(exc))
        db.rollback()
    finally:
        db.close()


# ── Core async generation logic ───────────────────────────────────────────────

async def async_generate_survey(
    request_id: str,
    data: Dict[str, Any],
    llm_model: str = "gpt",
) -> None:
    """
    Full survey generation pipeline:
      1. Fill missing planning context in the worker
      2. Generate questions via AI (JSON mode)
      3. Build DOCX document
      4. Upload to R2 (with local disk fallback)
      5. Build SurveyJS pages JSON
      6. Persist to DB as COMPLETED
    """
    logger.info("survey_generation_started", request_id=request_id, llm_model=llm_model)
    t0 = time.time()

    ai_service = AIService(llm_model=llm_model)
    try:
        await ai_service.initialize()
        update_survey_status(request_id, "RUNNING")
        await publish_progress(request_id, "STARTED")

        company_name = data["company_name"]
        project_name = data["project_name"]
        industry = data.get("industry", "")
        use_web_search = data.get("use_web_search", False)
        business_overview = (data.get("business_overview") or "").strip()
        use_case = (data.get("use_case") or "").strip()
        research_objectives = (data.get("research_objectives") or "").strip()

        if not business_overview:
            await publish_progress(request_id, "GENERATING_BUSINESS_OVERVIEW")
            business_overview = await ai_service.generate_business_overview(
                company_name=company_name,
                raw_input=use_case or industry or company_name,
            )
            update_survey_status(
                request_id,
                "RUNNING",
                business_overview=business_overview,
            )

        if not use_case:
            await publish_progress(request_id, "GENERATING_USE_CASE")
            use_case = await ai_service.generate_use_case(
                company_name=company_name,
                business_overview=business_overview,
            )
            update_survey_status(request_id, "RUNNING", use_case=use_case)

        if not research_objectives:
            await publish_progress(request_id, "GENERATING_RESEARCH_OBJECTIVES")
            research_objectives = await ai_service.generate_research_objectives(
                business_overview=business_overview,
                use_case=use_case,
            )
            update_survey_status(
                request_id,
                "RUNNING",
                research_objectives=research_objectives,
            )

        if not business_overview or not use_case or not research_objectives:
            raise ValueError("AI planning pipeline returned incomplete survey context")

        # ── 1. Generate questions ─────────────────────────────────────────────
        await publish_progress(request_id, "GENERATING_QUESTIONS")
        questions_json = await ai_service.generate_survey_questions(
            company_name=company_name,
            business_overview=business_overview,
            research_objectives=research_objectives,
            use_web_search=use_web_search,
        )
        logger.info("questions_generated", request_id=request_id)

        questions = _parse_questions_payload(questions_json)

        # ── 2. Build DOCX ─────────────────────────────────────────────────────
        await publish_progress(request_id, "BUILDING_DOCUMENT")
        assets_dir = Path(__file__).resolve().parent.parent / "assets"
        template_path = assets_dir / "template_new.docx"
        doc = Document(str(template_path)) if template_path.exists() else Document()

        doc.add_heading(f"{project_name} — Survey Questionnaire", 0)
        doc.add_paragraph(f"Company: {company_name}")
        doc.add_paragraph()

        for i, q in enumerate(questions.get("questions", []), 1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(q.get("text", "")).bold = True
            for choice in q.get("choices", []):
                doc.add_paragraph(str(choice), style="List Bullet 2")
            doc.add_paragraph()

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # ── 3. Upload to R2 with local fallback ───────────────────────────────
        await publish_progress(request_id, "UPLOADING_DOCUMENT")
        filename = f"{project_name.replace(' ', '_')}_questionnaire_{request_id}.docx"
        doc_link: str | None = None

        try:
            from app.services.storage_service import StorageService
            storage_service = StorageService()
            r2_url = await asyncio.to_thread(
                storage_service.upload_fileobj, doc_io, f"questionnaires/{filename}"
            )

            if r2_url:
                doc_link = f"/api/v1/files/download/{filename}"
                logger.info("r2_upload_success", request_id=request_id)
            else:
                logger.warning("r2_upload_returned_none", request_id=request_id)

        except Exception as exc:
            logger.warning("r2_upload_exception", request_id=request_id, error=str(exc))

        # Local fallback — write to disk if R2 failed or returned None
        if not doc_link:
            local_dir = Path(__file__).resolve().parent.parent / "questionnaires"
            local_dir.mkdir(parents=True, exist_ok=True)
            local_path = local_dir / filename
            doc_io.seek(0)
            local_path.write_bytes(doc_io.read())
            doc_link = f"/api/v1/files/download/{filename}"
            logger.warning(
                "r2_upload_failed_saved_locally",
                request_id=request_id,
                local_path=str(local_path),
            )

        # ── 4. Build SurveyJS pages ───────────────────────────────────────────
        pages = [{
            "name": "page1",
            "elements": [
                {
                    "type": q.get("type", "radiogroup"),
                    "name": f"q{i}",
                    "title": q.get("text", ""),
                    "choices": q.get("choices", []),
                }
                for i, q in enumerate(questions.get("questions", []), 1)
            ],
        }]

        # ── 5. Persist as COMPLETED ───────────────────────────────────────────
        update_survey_status(
            request_id,
            "COMPLETED",
            pages=pages,
            questionnaire_data=questions,
            doc_link=doc_link,
            business_overview=business_overview,
            use_case=use_case,
            research_objectives=research_objectives,
        )

        elapsed = round(time.time() - t0, 2)
        logger.info(
            "celery_task_completed_successfully",
            request_id=request_id,
            elapsed_seconds=elapsed,
        )
        await publish_progress(request_id, "SUCCESS")

    except Exception as exc:
        logger.error("survey_generation_failed", request_id=request_id, error=str(exc))
        update_survey_status(request_id, "FAILED")
        await publish_progress(request_id, f"ERROR: {exc}")
        raise

    finally:
        await ai_service.close()

@celery_app.task(
    bind=True,
    name="tasks.generate_survey",
    # REMOVED: autoretry_for, retry_backoff, retry_backoff_max, retry_jitter
    # Reason: having both autoretry_for AND manual self.retry() causes double-firing.
    # Manual retry below gives full control over countdown and state transitions.
    max_retries=3,
    default_retry_delay=60,
)
def generate_survey_task(
    self,
    request_id: str,
    data: Dict[str, Any],
    llm_model: str = "gpt",
) -> None:
    try:
        asyncio.run(async_generate_survey(request_id, data, llm_model))
    except Exception as exc:
        update_survey_status(request_id, "FAILED")

        retry_num = self.request.retries
        if retry_num < self.max_retries:
            countdown = 60 * (2 ** retry_num)   # 60s, 120s, 240s
            logger.warning(
                "celery_task_retrying",
                request_id=request_id,
                attempt=retry_num + 1,
                countdown_seconds=countdown,
                error=str(exc),
            )
            raise self.retry(exc=exc, countdown=countdown)

        logger.error(
            "celery_task_exhausted_retries",
            request_id=request_id,
            attempts=self.max_retries,
            error=str(exc),
        )
        raise


