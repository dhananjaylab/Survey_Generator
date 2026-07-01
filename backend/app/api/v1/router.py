"""
Survey generation API router.

P2 fix vs current repo state:
  list_surveys() used .all() with no cap — restored .limit(50) to prevent
  a full table scan / huge payload for users with many historical surveys.
  Everything else (Phase 2 Redis injection, Phase 3 _survey_access_filter,
  PUT /{id}/settings) is unchanged from the validated repo version.
"""
from fastapi import APIRouter, Depends, HTTPException, Request, status
from fastapi.responses import StreamingResponse
from io import BytesIO
from pathlib import Path
from urllib.parse import quote
from docx import Document
import asyncio
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
from sqlalchemy import or_
from typing import Optional

from app.models.database import get_db
from app.models.survey import SurveyRequestRecord
from app.services.ai_service import AIService
from app.core.auth import verify_token, get_current_user
from app.core.rate_limit import limiter
from app.core.logging import get_logger
from app.core.metrics import get_metrics_collector

logger  = get_logger(__name__)
metrics = get_metrics_collector()

router = APIRouter(
    prefix="/api/v1/surveys",
    tags=["Surveys"],
    dependencies=[Depends(verify_token)],
)


# ── Ownership filter ──────────────────────────────────────────────────────────

def _survey_access_filter(current_user: str):
    """
    Legacy surveys created before username ownership was stored may have a NULL
    username. Keep them visible to the signed-in user so older history does not
    disappear after auth migrations.
    """
    return or_(
        SurveyRequestRecord.username == current_user,
        SurveyRequestRecord.username.is_(None),
    )


# ── Redis dependency ──────────────────────────────────────────────────────────

def get_redis(request: Request):
    """Inject the shared Redis pool from app state (created at startup)."""
    return request.app.state.redis


def _build_docx_from_pages(project_name: str, company_name: str, pages: list) -> BytesIO:
    assets_dir = Path(__file__).resolve().parent.parent / "assets"
    template_path = assets_dir / "template_new.docx"
    doc = Document(str(template_path)) if template_path.exists() else Document()

    doc.add_heading(f"{project_name} — Survey Questionnaire", 0)
    doc.add_paragraph(f"Company: {company_name}")
    doc.add_paragraph()

    for page in pages or []:
        for element in page.get("elements", []):
            p = doc.add_paragraph(style="List Number")
            p.add_run(element.get("title", "")).bold = True
            for choice in element.get("choices", []):
                if isinstance(choice, dict):
                    choice_text = choice.get("text") or choice.get("value") or ""
                else:
                    choice_text = str(choice)
                if choice_text:
                    doc.add_paragraph(choice_text, style="List Bullet 2")
            doc.add_paragraph()

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ── Request / Response models ─────────────────────────────────────────────────

class BusinessOverviewRequest(BaseModel):
    company_name: str = Field(..., min_length=1, max_length=200)
    raw_input:    str = Field(..., min_length=1, max_length=2000)
    llm_model:    str = Field(default="gpt")


class UseCaseRequest(BaseModel):
    company_name:      str = Field(..., min_length=1, max_length=200)
    business_overview: str = Field(..., min_length=1, max_length=5000)
    llm_model:         str = Field(default="gpt")


class ResearchObjectivesRequest(BaseModel):
    business_overview: str = Field(..., min_length=10, max_length=5000)
    use_case:          str = Field(..., min_length=10, max_length=2000)
    llm_model:         str = Field(default="gpt")


class GenerateSurveyRequest(BaseModel):
    request_id:          str = Field(...)
    project_name:        str = Field(..., min_length=1, max_length=200)
    company_name:        str = Field(..., min_length=1, max_length=200)
    business_overview:   str = Field(default="", max_length=5000)
    research_objectives: str = Field(default="", max_length=2000)
    industry:            str = Field(..., max_length=100)
    use_case:            str = Field(default="", max_length=2000)
    llm_model:           str = Field(default="gpt")
    use_web_search:      bool = Field(default=False)


class RegenerateDocumentRequest(BaseModel):
    request_id:         str
    project_name:       str = Field(..., min_length=1, max_length=200)
    company_name:       str = Field(default="", max_length=200)
    survey_title:       str = Field(default="", max_length=200)
    # This is optional metadata for future export templates; keep it generous
    # so long business summaries do not trip FastAPI validation.
    survey_description: str = Field(default="", max_length=5000)
    pages:              list = Field(default_factory=list)
    delivery_mode:      str = Field(default="none")  # none | local | r2


class SurveySettingsRequest(BaseModel):
    llm_model:      Optional[str] = None
    use_web_search: Optional[bool] = None


# ── AI endpoints (shared Redis cache) ─────────────────────────────────────────

@router.post("/business-overview")
@limiter.limit("10/minute")
async def get_business_overview(
    request: Request,
    req: BusinessOverviewRequest,
    redis=Depends(get_redis),
    current_user: str = Depends(get_current_user),
):
    logger.info("business_overview_requested", user=current_user)
    service = AIService(llm_model=req.llm_model, redis=redis)
    await service.initialize()
    try:
        result = await service.generate_business_overview(
            company_name=req.company_name,
            raw_input=req.raw_input,
        )
        return {"business_overview": result}
    finally:
        await service.close()


@router.post("/generate-use-case")
@limiter.limit("10/minute")
async def generate_use_case(
    request: Request,
    req: UseCaseRequest,
    redis=Depends(get_redis),
    current_user: str = Depends(get_current_user),
):
    logger.info("use_case_requested", user=current_user)
    service = AIService(llm_model=req.llm_model, redis=redis)
    await service.initialize()
    try:
        result = await service.generate_use_case(
            company_name=req.company_name,
            business_overview=req.business_overview,
        )
        return {"use_case": result}
    finally:
        await service.close()


@router.post("/research-objectives")
@limiter.limit("10/minute")
async def get_research_objectives(
    request: Request,
    req: ResearchObjectivesRequest,
    redis=Depends(get_redis),
    current_user: str = Depends(get_current_user),
):
    logger.info("research_objectives_requested", user=current_user)
    service = AIService(llm_model=req.llm_model, redis=redis)
    await service.initialize()
    try:
        result = await service.generate_research_objectives(
            business_overview=req.business_overview,
            use_case=req.use_case,
        )
        return {"research_objectives": result}
    finally:
        await service.close()


# ── Survey CRUD + generation ──────────────────────────────────────────────────

@router.post("/generate", status_code=status.HTTP_202_ACCEPTED)
@limiter.limit("5/minute")
async def generate_survey(
    request: Request,
    req: GenerateSurveyRequest,
    db: Session = Depends(get_db),
    current_user: str = Depends(get_current_user),
):
    logger.info("survey_generate_requested", request_id=req.request_id, user=current_user)
    metrics.record_survey_started()

    record = SurveyRequestRecord(
        request_id=req.request_id,
        username=current_user,
        project_name=req.project_name,
        company_name=req.company_name,
        industry=req.industry,
        use_case=req.use_case or None,
        business_overview=req.business_overview or None,
        research_objectives=req.research_objectives or None,
        status="PENDING",
    )
    db.add(record)
    db.commit()

    from app.tasks.survey_tasks import generate_survey_task
    generate_survey_task.delay(
        request_id=req.request_id,
        data=req.model_dump() | {"delivery_mode": "none"},
        llm_model=req.llm_model,
    )

    return {
        "request_id": req.request_id,
        "status": "PENDING",
        "message": "Survey generation started",
    }


@router.get("/status/{request_id}")
async def get_survey_status(
    request_id: str,
    current_user: str = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    record = (
        db.query(SurveyRequestRecord)
        .filter(
            SurveyRequestRecord.request_id == request_id,
            _survey_access_filter(current_user),
        )
        .first()
    )
    if not record:
        raise HTTPException(status_code=404, detail="Survey not found")

    return {
        "request_id":          record.request_id,
        "status":              record.status,
        "doc_link":            record.doc_link,
        "pages":               record.pages,
        "project_name":        record.project_name,
        "company_name":        record.company_name,
        "industry":            record.industry,
        "use_case":            record.use_case,
        "business_overview":   record.business_overview,
        "research_objectives": record.research_objectives,
    }


@router.get("/")
async def list_surveys(
    current_user: str = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    records = (
        db.query(SurveyRequestRecord)
        .filter(_survey_access_filter(current_user))
        .order_by(SurveyRequestRecord.created_at.desc())
        .limit(50)   # P2 fix: was unbounded .all() in repo
        .all()
    )
    return {
        "surveys": [
            {
                "request_id":   r.request_id,
                "project_name": r.project_name,
                "company_name": r.company_name,
                "status":       r.status,
                "doc_link":     r.doc_link,
                "created_at":   r.created_at.isoformat() if r.created_at else None,
            }
            for r in records
        ],
        "success": 1,
    }


@router.delete("/{request_id}")
async def delete_survey(
    request_id: str,
    current_user: str = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    record = (
        db.query(SurveyRequestRecord)
        .filter(
            SurveyRequestRecord.request_id == request_id,
            _survey_access_filter(current_user),
        )
        .first()
    )
    if not record:
        raise HTTPException(status_code=404, detail="Survey not found")

    db.delete(record)
    db.commit()
    logger.info("survey_deleted", request_id=request_id, user=current_user)
    return {"message": "Survey deleted", "success": 1}


@router.post("/export-local")
@limiter.limit("5/minute")
async def export_survey_local(
    request: Request,
    req: RegenerateDocumentRequest,
    current_user: str = Depends(get_current_user),
):
    doc_io = _build_docx_from_pages(req.project_name, req.company_name, req.pages)
    filename = f"{req.project_name.replace(' ', '_')}_survey.docx"
    return StreamingResponse(
        doc_io,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/export-document")
@limiter.limit("5/minute")
async def export_survey_document(
    request: Request,
    req: RegenerateDocumentRequest,
    current_user: str = Depends(get_current_user),
):
    doc_io = _build_docx_from_pages(req.project_name, req.company_name, req.pages)
    filename = f"{req.project_name.replace(' ', '_')}_survey.docx"
    local_link = f"/api/v1/files/download/{quote(filename)}"

    delivery_mode = (req.delivery_mode or "none").strip().lower()
    if delivery_mode not in {"local", "r2"}:
        raise HTTPException(status_code=400, detail="delivery_mode must be local or r2")

    doc_link: str | None = None
    if delivery_mode == "local":
        questionnaires_dir = Path(__file__).resolve().parent.parent.parent / "questionnaires"
        questionnaires_dir.mkdir(parents=True, exist_ok=True)
        (questionnaires_dir / filename).write_bytes(doc_io.getvalue())
        doc_link = local_link

    elif delivery_mode == "r2":
        logger.info("export_r2_requested", request_id=req.request_id, bucket_mode=delivery_mode, filename=filename)
        from app.services.storage_service import StorageService
        storage_service = StorageService()
        r2_url = await asyncio.to_thread(storage_service.upload_fileobj, doc_io, f"questionnaires/{filename}")
        if not r2_url:
            raise HTTPException(status_code=500, detail="R2 upload failed")
        doc_link = r2_url

    return {"success": 1, "message": "Document exported", "doc_link": doc_link, "request_id": req.request_id}


@router.put("/{survey_id}/settings")
async def update_survey_settings(
    survey_id: str,
    settings: dict,
    current_user: str = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Persist behavioural trigger settings from the Builder page."""
    record = (
        db.query(SurveyRequestRecord)
        .filter(
            SurveyRequestRecord.request_id == survey_id,
            _survey_access_filter(current_user),
        )
        .first()
    )
    if not record:
        raise HTTPException(status_code=404, detail="Survey not found")

    record.settings = settings
    db.commit()
    return {"message": "Settings updated", "success": 1}


@router.get("/settings")
async def get_settings(current_user: str = Depends(get_current_user)):
    return {
        "available_models": ["gpt", "gemini"],
        "default_model":    "gpt",
        "web_search":       False,
    }
